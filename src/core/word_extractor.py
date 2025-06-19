"""
Word数据提取模块
从Word文档中提取数据
"""
import os
import re
import sys
import pandas as pd
from docx import Document
from loguru import logger

from src.core.placeholder_parser import PlaceholderParserFactory, BasePlaceholderParser
from itertools import zip_longest


class WordExtractor:
    """
    Word数据提取器，从Word文档中提取数据
    """
    
    def __init__(self, config_manager=None):
        """
        初始化Word数据提取器
        
        Args:
            config_manager: 配置管理器，可选
        """
        logger.info("初始化Word数据提取器")
        
        # 保存配置管理器
        self.config_manager = config_manager
        
        # 使用解析器工厂
        self.placeholder_parser_factory = PlaceholderParserFactory()
        
        # 定义支持的占位符模式
        self.placeholder_patterns = [
            r'\{\{(.*?)\}\}',  # {{变量}}
            r'\$\{(.*?)\}',    # ${变量}
            r'#(.*?)#',        # #变量#
            r'\{\{img:(.*?)\}\}'  # {{img:变量}}
        ]
        
        # 模板文档
        self.template_doc = None
    
    def detect_placeholders(self, doc_path_or_doc):
        """
        检测文档中的所有占位符
        
        Args:
            doc_path_or_doc: 文档路径或Document对象
            
        Returns:
            list: 占位符列表
        """
        try:
            # 如果传入的是路径，则加载文档
            if isinstance(doc_path_or_doc, str):
                doc = Document(doc_path_or_doc)
            else:
                doc = doc_path_or_doc
                
            # 提取所有文本
            all_text = ""
            
            # 处理段落
            for para in doc.paragraphs:
                all_text += para.text + "\n"
                
            # 处理表格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        all_text += cell.text + "\n"
            
            # 使用解析器工厂查找所有占位符
            placeholders = self.placeholder_parser_factory.find_all_placeholders(all_text)
            
            logger.info(f"从模板中检测到{len(placeholders)}个占位符")
            return placeholders
            
        except Exception as e:
            logger.error(f"检测占位符时出错: {str(e)}")
            return []
    
    def extract_data(self, target_doc_path, template_doc=None):
        """
        从目标文档中提取数据
        
        Args:
            target_doc_path: 目标文档路径
            template_doc: 模板文档对象，如果为None则使用已加载的模板
            
        Returns:
            dict: 提取的数据，格式为 {placeholder: value}
        """
        try:
            if not template_doc and not self.template_doc:
                raise ValueError("未提供模板文档")
                
            template = template_doc if template_doc else self.template_doc
            
            # 设置当前处理的文档路径（用于图片提取）
            self.current_doc_path = target_doc_path
            
            # 打开目标文档
            target_doc = Document(target_doc_path)
            
            # 初始化结果字典
            extracted_data = {}
            
            # 提取占位符列表
            placeholders = self.detect_placeholders(template)
            if not placeholders:
                logger.warning("未在模板中找到任何占位符")
                return {}
            
            # 先处理非图片占位符
            non_image_placeholders = [p for p in placeholders if not (p.startswith("img:") or "img:" in p)]
            for placeholder in non_image_placeholders:
                # 根据模板中的占位符位置进行提取
                value = ""
                
                # 按照段落查找
                for t_para, d_para in zip_longest(template.paragraphs, target_doc.paragraphs):
                    if t_para and d_para:
                        # 如果段落中有当前占位符，从对应目标段落提取值
                        t_text = t_para.text
                        d_text = d_para.text
                        
                        # 检查当前占位符是否存在于模板段落中
                        parser = self.placeholder_parser_factory.get_parser_for_text(t_text)
                        if parser and parser.has_placeholder(t_text, placeholder):
                            # 提取值
                            extracted_value = self._extract_value_for_placeholder(placeholder, t_text, d_text)
                            if extracted_value:
                                value = extracted_value
                                break
                
                # 如果段落中未找到，则检查表格
                if not value:
                    value = self._extract_from_tables(template, target_doc, placeholder)
                
                # 保存提取的值
                if value:
                    extracted_data[placeholder] = value
            
            # 保存当前已提取的数据，供图片命名使用
            self.current_extracted_data = extracted_data
            
            # 处理图片占位符
            image_placeholders = [p for p in placeholders if p.startswith("img:") or "img:" in p]
            for placeholder in image_placeholders:
                # 根据模板中的占位符位置进行提取
                value = ""
                
                # 按照段落查找
                for t_para, d_para in zip_longest(template.paragraphs, target_doc.paragraphs):
                    if t_para and d_para:
                        # 如果段落中有当前占位符，从对应目标段落提取值
                        t_text = t_para.text
                        d_text = d_para.text
                        
                        # 检查当前占位符是否存在于模板段落中
                        parser = self.placeholder_parser_factory.get_parser_for_text(t_text)
                        if parser and parser.has_placeholder(t_text, placeholder):
                            # 提取值
                            extracted_value = self._extract_value_for_placeholder(placeholder, t_text, d_text)
                            if extracted_value:
                                value = extracted_value
                                break
                
                # 如果段落中未找到，则检查表格
                if not value:
                    value = self._extract_from_tables(template, target_doc, placeholder)
                
                # 保存提取的值
                if value:
                    extracted_data[placeholder] = value
            
            # 记录提取的数据行数
            if extracted_data:
                logger.info(f"成功从目标文档中提取数据，共{len(extracted_data)}行")
            else:
                logger.warning(f"未从文档 {target_doc_path} 中提取到数据")
                
            return extracted_data
            
        except Exception as e:
            logger.error(f"提取数据时出错: {str(e)}")
            raise
    
    def _extract_from_tables(self, template_doc, target_doc, placeholder):
        """
        从表格中提取占位符对应的值
        
        Args:
            template_doc: 模板文档对象
            target_doc: 目标文档对象
            placeholder: 要查找的占位符名称
            
        Returns:
            str: 提取的值，如果未找到则返回空字符串
        """
        if not hasattr(template_doc, 'tables') or not hasattr(target_doc, 'tables'):
            return ""
            
        # 遍历所有表格
        for t_idx, t_table in enumerate(template_doc.tables):
            if t_idx >= len(target_doc.tables):
                break
                
            d_table = target_doc.tables[t_idx]
            
            # 创建占位符位置映射
            placeholder_positions = {}
            date_placeholders_in_row = {}
            
            # 首先扫描表格，记录所有占位符的位置
            for r_idx, t_row in enumerate(t_table.rows):
                if r_idx >= len(d_table.rows):
                    break
                    
                # 检查此行是否包含日期相关占位符
                has_start_date = False
                has_end_date = False
                start_date_col = -1
                end_date_col = -1
                
                for c_idx, t_cell in enumerate(t_row.cells):
                    if c_idx >= len(d_table.rows[r_idx].cells):
                        break
                        
                    t_text = t_cell.text
                    
                    # 记录占位符位置
                    for pattern in self.placeholder_patterns:
                        matches = re.findall(pattern, t_text)
                        for match in matches:
                            if isinstance(match, tuple):  # 处理正则表达式组
                                match = match[0]
                            if match == placeholder:
                                placeholder_positions[(r_idx, c_idx)] = placeholder
                            
                            # 特殊处理日期占位符位置
                            if match == "开始日期":
                                has_start_date = True
                                start_date_col = c_idx
                            elif match == "结束日期":
                                has_end_date = True
                                end_date_col = c_idx
                
                # 记录同一行中的日期占位符
                if has_start_date and has_end_date:
                    date_placeholders_in_row[r_idx] = (start_date_col, end_date_col)
            
            # 遍历表格的每个单元格
            for r_idx, t_row in enumerate(t_table.rows):
                if r_idx >= len(d_table.rows):
                    break
                    
                d_row = d_table.rows[r_idx]
                
                # 特殊处理：如果当前行包含日期范围占位符，且当前查找的是其中之一
                if r_idx in date_placeholders_in_row and (placeholder == "开始日期" or placeholder == "结束日期"):
                    start_col, end_col = date_placeholders_in_row[r_idx]
                    
                    # 如果查找的是开始日期
                    if placeholder == "开始日期" and start_col >= 0 and start_col < len(d_row.cells):
                        d_cell = d_row.cells[start_col]
                        d_text = d_cell.text.strip()
                        
                        # 检查是否是日期范围格式的文本
                        date_range_pattern = r'(\d{4}年\s*\d{1,2}月\s*\d{1,2}日)\s*[至到\-~]\s*(\d{4}年\s*\d{1,2}月\s*\d{1,2}日)'
                        date_range_match = re.search(date_range_pattern, d_text)
                        
                        if date_range_match:
                            return date_range_match.group(1).strip()
                        elif d_text and "至" not in d_text and "到" not in d_text:
                            # 如果没有日期范围，但有单独的日期，直接返回
                            return d_text
                    
                    # 如果查找的是结束日期
                    elif placeholder == "结束日期" and end_col >= 0 and end_col < len(d_row.cells):
                        d_cell = d_row.cells[end_col]
                        d_text = d_cell.text.strip()
                        
                        # 检查是否是日期范围格式的文本
                        date_range_pattern = r'(\d{4}年\s*\d{1,2}月\s*\d{1,2}日)\s*[至到\-~]\s*(\d{4}年\s*\d{1,2}月\s*\d{1,2}日)'
                        date_range_match = re.search(date_range_pattern, d_text)
                        
                        if date_range_match:
                            return date_range_match.group(2).strip()
                        elif d_text and "至" not in d_text and "到" not in d_text:
                            # 如果没有日期范围，但有单独的日期，直接返回
                            return d_text
                            
                    # 特殊情况：开始和结束日期在同一个单元格
                    if start_col == end_col:
                        d_cell = d_row.cells[start_col]
                        d_text = d_cell.text.strip()
                        
                        # 检查是否是日期范围格式
                        date_range_pattern = r'(\d{4}年\s*\d{1,2}月\s*\d{1,2}日)\s*[至到\-~]\s*(\d{4}年\s*\d{1,2}月\s*\d{1,2}日)'
                        date_range_match = re.search(date_range_pattern, d_text)
                        
                        if date_range_match:
                            if placeholder == "开始日期":
                                return date_range_match.group(1).strip()
                            else:  # 结束日期
                                return date_range_match.group(2).strip()
                
                # 正常处理单个占位符
                for c_idx, t_cell in enumerate(t_row.cells):
                    if c_idx >= len(d_row.cells):
                        break
                        
                    d_cell = d_row.cells[c_idx]
                    
                    # 获取单元格文本
                    t_text = t_cell.text
                    d_text = d_cell.text
                    
                    # 检查占位符是否存在于模板单元格中
                    parser = self.placeholder_parser_factory.get_parser_for_text(t_text)
                    if parser and parser.has_placeholder(t_text, placeholder):
                        # 提取值
                        value = self._extract_value_for_placeholder(placeholder, t_text, d_text)
                        if value:
                            return value
                            
        # 特殊处理：查找带有标签行的表格数据
        for t_idx, t_table in enumerate(template_doc.tables):
            if t_idx >= len(target_doc.tables):
                break
                
            d_table = target_doc.tables[t_idx]
            
            # 查找带有占位符的标题行
            header_row_idx = -1
            header_col_idx = -1
            
            for r_idx, row in enumerate(t_table.rows):
                for c_idx, cell in enumerate(row.cells):
                    parser = self.placeholder_parser_factory.get_parser_for_text(cell.text)
                    if parser and parser.has_placeholder(cell.text, placeholder):
                        header_row_idx = r_idx
                        header_col_idx = c_idx
                        break
                if header_row_idx >= 0:
                    break
            
            # 如果找到了标题行，查找对应列中的数据
            if header_row_idx >= 0 and header_col_idx >= 0:
                # 查找与占位符匹配的标签文本
                label_text = None
                for r_idx in range(header_row_idx + 1, len(t_table.rows)):
                    if r_idx < len(d_table.rows):
                        cell_text = t_table.rows[r_idx].cells[header_col_idx].text
                        if cell_text and not any(parser.has_placeholder(cell_text, p) for p in self.detect_placeholders(template_doc)):
                            label_text = cell_text
                            break
                
                # 使用找到的标签文本在目标表格中查找对应的数据
                if label_text:
                    for r_idx in range(len(d_table.rows)):
                        for c_idx in range(len(d_table.rows[r_idx].cells)):
                            if d_table.rows[r_idx].cells[c_idx].text == label_text and c_idx + 1 < len(d_table.rows[r_idx].cells):
                                return d_table.rows[r_idx].cells[c_idx + 1].text.strip()
        
        # 处理表格内相邻单元格的情况
        for t_idx, t_table in enumerate(template_doc.tables):
            if t_idx >= len(target_doc.tables):
                break
                
            d_table = target_doc.tables[t_idx]
            
            for r_idx, t_row in enumerate(t_table.rows):
                if r_idx >= len(d_table.rows):
                    break
                
                d_row = d_table.rows[r_idx]
                
                for c_idx, t_cell in enumerate(t_row.cells):
                    if c_idx >= len(d_row.cells) - 1:
                        break
                    
                    t_text = t_cell.text
                    # 检查当前单元格是否包含占位符名称（不带括号等格式）
                    if t_text.strip() == placeholder:
                        # 返回相邻单元格的值
                        return d_row.cells[c_idx + 1].text.strip()
        
        return ""
    
    def _extract_table_cell_value(self, placeholder, template_text, target_text):
        """
        从表格单元格中提取占位符对应的值
        
        Args:
            placeholder: 占位符名称
            template_text: 模板单元格文本
            target_text: 目标单元格文本
            
        Returns:
            str: 提取的值，如果未找到则返回空字符串
        """
        # 查找对应的占位符格式
        placeholder_matches = []
        for pattern in self.placeholder_patterns:
            placeholder_pattern = pattern.replace("(.*?)", re.escape(placeholder))
            if re.search(placeholder_pattern, template_text):
                placeholder_matches.append((re.escape(placeholder_pattern), placeholder_pattern))
        
        if not placeholder_matches:
            return ""
        
        for escaped_pattern, raw_pattern in placeholder_matches:
            # 提取方法1: 标签前置的值
            # 例如: "姓名：{{姓名}}" -> "姓名：张三"
            label_pattern = r'([^:：]+)[:：]\s*' + escaped_pattern
            label_match = re.search(label_pattern, template_text)
            if label_match:
                label = label_match.group(1).strip()
                # 在目标文本中查找相同标签后面的值
                value_pattern = re.escape(label) + r'[:：]\s*(.*?)(?=$|\n|,|，|;|；|\s\s)'
                value_match = re.search(value_pattern, target_text)
                if value_match:
                    return value_match.group(1).strip()
            
            # 提取方法2: 位置替换
            # 例如: "姓名：{{姓名}} 性别：" -> "姓名：张三 性别："
            parts = re.split(escaped_pattern, template_text)
            if len(parts) == 2:
                prefix = parts[0]
                suffix = parts[1]
                
                # 如果前缀和后缀都存在于目标文本中
                if prefix in target_text and suffix in target_text:
                    # 提取前缀和后缀之间的文本作为值
                    pattern = re.escape(prefix) + r'(.*?)' + re.escape(suffix)
                    match = re.search(pattern, target_text)
                    if match:
                        return match.group(1).strip()
        
        # 提取方法3: 单独值
        # 如果目标单元格只包含一个简单值（没有复杂格式、占位符或标签）
        target_text = target_text.strip()
        if target_text and len(target_text.split()) <= 3:
            # 确保目标文本不包含占位符格式
            contains_placeholder = False
            for pattern in self.placeholder_patterns:
                placeholder_pattern = pattern.replace("(.*?)", r"[^{}]+?")
                if re.search(placeholder_pattern, target_text):
                    contains_placeholder = True
                    break
            
            # 确保不是标签形式
            if not contains_placeholder and not re.search(r'[:：]', target_text):
                return target_text
        
        return ""
    
    def _extract_from_table_headers(self, template_doc, target_doc, placeholders, extracted_values):
        """
        从表格表头提取数据
        
        Args:
            template_doc: 模板文档
            target_doc: 目标文档
            placeholders: 占位符列表
            extracted_values: 存储提取值的字典
        """
        # 遍历模板文档中的所有表格
        for table_idx, template_table in enumerate(template_doc.tables):
            if table_idx >= len(target_doc.tables):
                continue
            
            target_table = target_doc.tables[table_idx]
            headers_map = {}  # 存储表头与列索引的映射
            
            # 检查第一行是否为表头
            if len(template_table.rows) > 0 and len(target_table.rows) > 0:
                template_header_row = template_table.rows[0]
                target_header_row = target_table.rows[0]
                
                # 识别表头中的占位符
                for cell_idx, cell in enumerate(template_header_row.cells):
                    if cell_idx >= len(target_header_row.cells):
                        continue
                    
                    cell_text = cell.text.strip()
                    
                    # 检查单元格是否包含占位符
                    for placeholder in placeholders:
                        for pattern in self.placeholder_patterns:
                            placeholder_pattern = pattern.replace("(.*?)", re.escape(placeholder))
                            if re.search(placeholder_pattern, cell_text):
                                # 找到了表头中的占位符，记录列索引
                                headers_map[placeholder] = cell_idx
                
                # 如果在表头中找到了占位符，检查数据行
                if headers_map and len(template_table.rows) > 1 and len(target_table.rows) > 1:
                    # 假设第二行包含实际数据
                    data_row = target_table.rows[1]
                    
                    # 从数据行中提取值
                    for placeholder, col_idx in headers_map.items():
                        if col_idx < len(data_row.cells):
                            value = data_row.cells[col_idx].text.strip()
                            
                            if value and placeholder not in extracted_values:
                                extracted_values[placeholder] = value
                                logger.debug(f"从表格表头映射提取到 {placeholder}: {value}")
        
        return extracted_values
    
    def _extract_from_tables_fallback(self, doc, data, placeholders):
        """
        从表格中提取数据的备用方法
        
        Args:
            doc: Word文档
            data: 数据存储字典
            placeholders: 占位符列表
        """
        # 从表格中查找可能的数据
        for table in doc.tables:
            # 检查每一行
            for row in table.rows:
                # 获取这一行的所有文本
                row_text = " ".join([cell.text for cell in row.cells])
                
                # 检查这一行是否包含任何占位符的标签部分
                for placeholder in placeholders:
                    # 先检查是否有"标签: 值"格式
                    if placeholder + ":" in row_text or placeholder + "：" in row_text:
                        # 尝试提取值
                        pattern = re.escape(placeholder) + r'[:：]\s*(.*?)(?=$|\n|,|，|;|；|\s\s)'
                        matches = re.findall(pattern, row_text)
                        if matches:
                            value = matches[0].strip()
                            if value and placeholder in data:
                                data[placeholder].append(value)
                                logger.debug(f"从表格行文本中提取到 {placeholder}: {value}")
                
                # 逐个单元格检查
                for i, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    
                    # 检查单元格是否是标签
                    for placeholder in placeholders:
                        if cell_text == placeholder or placeholder.lower() == cell_text.lower():
                            # 检查下一个单元格是否包含值
                            if i + 1 < len(row.cells):
                                value = row.cells[i + 1].text.strip()
                                if value and placeholder in data:
                                    data[placeholder].append(value)
                                    logger.debug(f"从表格相邻单元格中提取到 {placeholder}: {value}")
                    
                    # 检查单元格是否包含"标签: 值"格式
                    for placeholder in placeholders:
                        label_value_pattern = re.escape(placeholder) + r'[:：]\s*(.*?)$'
                        matches = re.findall(label_value_pattern, cell_text)
                        if matches:
                            value = matches[0].strip()
                            if value and placeholder in data:
                                data[placeholder].append(value)
                                logger.debug(f"从表格单元格中提取到 {placeholder}: {value}")
                    
                    # 尝试通过占位符模式匹配提取值
                    for placeholder in placeholders:
                        for pattern in self.placeholder_patterns:
                            placeholder_text = pattern.replace("(.*?)", placeholder)
                            if placeholder_text in cell_text:
                                # 简单替换，去除占位符文本
                                value = cell_text.replace(placeholder_text, "").strip()
                                if value and placeholder in data:
                                    data[placeholder].append(value)
                                    logger.debug(f"通过替换提取到 {placeholder}: {value}")
    
    def extract_batch_data(self, template_path, target_dir, output_excel=None, recursive=False):
        """
        从目标目录中批量提取数据
        
        Args:
            template_path: 模板文件路径
            target_dir: 目标文件目录
            output_excel: 输出Excel文件路径，可选
            recursive: 是否递归处理子目录，默认为False
            
        Returns:
            pandas.DataFrame: 提取的数据
        """
        try:
            # 收集目标文件
            target_files = []
            
            if recursive:
                # 递归处理
                for root, dirs, files in os.walk(target_dir):
                    for file in files:
                        if file.endswith('.docx') and not file.startswith('~$'):
                            file_path = os.path.join(root, file)
                            if file_path != template_path:  # 排除模板文件
                                target_files.append(file_path)
            else:
                # 只处理当前目录
                for file in os.listdir(target_dir):
                    if file.endswith('.docx') and not file.startswith('~$'):
                        file_path = os.path.join(target_dir, file)
                        if file_path != template_path:  # 排除模板文件
                            target_files.append(file_path)
            
            # 调用文件列表处理方法
            logger.info(f"找到{len(target_files)}个Word文件进行处理")
            return self.extract_batch_data_from_files(template_path, target_files, output_excel)
            
        except Exception as e:
            logger.error(f"批量提取数据时出错: {str(e)}")
            raise
    
    def extract_batch_data_from_files(self, template_path, target_files, output_excel=None):
        """
        从多个目标文件中批量提取数据
        
        Args:
            template_path: 模板文件路径
            target_files: 目标文件路径列表
            output_excel: 输出Excel文件路径，可选
            
        Returns:
            pandas.DataFrame: 提取的数据
        """
        try:
            # 加载模板
            self.template_doc = Document(template_path)
            placeholders = self.detect_placeholders(self.template_doc)
            
            if not placeholders:
                logger.warning("未在模板中找到任何占位符")
                return pd.DataFrame()
            
            # 初始化数据列表和文件计数
            all_data = []
            file_count = 0
            
            # 处理每个目标文件
            for target_file in target_files:
                try:
                    logger.info(f"处理文件: {os.path.basename(target_file)}")
                    
                    # 重置当前提取的数据
                    self.current_extracted_data = {}
                    
                    # 设置当前处理的文档路径（用于图片提取）
                    self.current_doc_path = target_file
                    
                    # 提取数据
                    file_data = self.extract_data(target_file, self.template_doc)
                    
                    if file_data:
                        # 添加文件路径到数据中
                        file_data['文件路径'] = target_file
                        file_data['文件名'] = os.path.basename(target_file)
                        
                        # 添加到数据列表
                        all_data.append(file_data)
                        file_count += 1
                        logger.info(f"从 {os.path.basename(target_file)} 提取了{len(file_data)}个字段")
                    else:
                        logger.warning(f"从 {os.path.basename(target_file)} 未提取到数据")
                except Exception as e:
                    logger.error(f"处理文件 {os.path.basename(target_file)} 时出错: {str(e)}")
            
            # 转换为DataFrame
            if all_data:
                df = pd.DataFrame(all_data)
                
                # 特殊处理img:照片字段
                for col in df.columns:
                    if col.startswith("img:") or "img:" in col:
                        logger.info(f"发现图片字段: {col}")
                        # 确保图片路径存在
                        df[col] = df[col].apply(
                            lambda x: x if x and isinstance(x, str) and os.path.exists(x) else "")
                
                # 保存到Excel
                if output_excel:
                    try:
                        df.to_excel(output_excel, index=False)
                        logger.info(f"数据已保存到Excel文件: {output_excel}")
                    except Exception as e:
                        logger.error(f"保存Excel文件时出错: {str(e)}")
                
                logger.info(f"成功处理{file_count}个文件，提取了{len(df)}条数据")
                return df
            else:
                logger.warning("未从任何文件中提取到数据")
                return pd.DataFrame()
        except Exception as e:
            logger.error(f"批量提取数据时出错: {str(e)}")
            raise
    
    def _extract_value_for_placeholder(self, placeholder, template_text, target_text):
        """
        根据模板文本和目标文本提取占位符对应的值
        
        Args:
            placeholder: 占位符名称
            template_text: 模板文本
            target_text: 目标文本
            
        Returns:
            str: 提取的值，如果未找到则返回空字符串
        """
        # 特殊处理图片占位符
        if placeholder.startswith("img:") or "img:" in placeholder:
            try:
                # 尝试从文档中提取图片
                image_path = self._extract_image_from_document(placeholder)
                if image_path:
                    logger.info(f"成功提取图片: {image_path}")
                    return image_path
                else:
                    logger.warning(f"未能提取到图片: {placeholder}")
                    return "未能提取到图片"
            except Exception as e:
                logger.error(f"提取图片时出错: {str(e)}")
                return f"图片提取错误: {str(e)}"
            
        # 特殊处理日期格式
        if placeholder == "开始日期" or placeholder == "结束日期":
            # 查找日期范围格式的文本
            date_range_patterns = [
                r'(\d{4}年\s*\d{1,2}月\s*\d{1,2}日)\s*[至到\-~]\s*(\d{4}年\s*\d{1,2}月\s*\d{1,2}日)',  # 完整日期范围
                r'(\d{4}年\s*\d{1,2}月\s*\d{1,2}日)',  # 单独日期
                r'(\d{4}[-/]\d{1,2}[-/]\d{1,2})\s*[至到\-~]\s*(\d{4}[-/]\d{1,2}[-/]\d{1,2})',  # 数字格式日期范围
                r'(\d{4}[-/]\d{1,2}[-/]\d{1,2})',  # 数字格式单独日期
                r'(年\s+月\s+日)\s+[至到\-~]\s*\n?(年\s+月\s+日)'  # 未填写完整的日期范围
            ]
            
            for pattern in date_range_patterns:
                date_match = re.search(pattern, target_text)
                if date_match:
                    if len(date_match.groups()) == 2:  # 日期范围
                        if placeholder == "开始日期":
                            return date_match.group(1).strip()
                        else:  # 结束日期
                            return date_match.group(2).strip()
                    else:  # 单独日期
                        return date_match.group(1).strip()
            
            # 如果没有找到匹配的日期格式，尝试从文本中提取可能的日期
            # 查找形如"YYYY年MM月DD日"的文本
            single_date_pattern = r'\d{4}年\s*\d{1,2}月\s*\d{1,2}日'
            single_date_match = re.search(single_date_pattern, target_text)
            if single_date_match:
                return single_date_match.group(0).strip()
            
            # 查找形如"YYYY-MM-DD"或"YYYY/MM/DD"的文本
            numeric_date_pattern = r'\d{4}[-/]\d{1,2}[-/]\d{1,2}'
            numeric_date_match = re.search(numeric_date_pattern, target_text)
            if numeric_date_match:
                return numeric_date_match.group(0).strip()
        
        # 在模板文本中找到完整的占位符表示
        placeholder_matches = []
        standard_pattern = r'\{\{' + re.escape(placeholder) + r'\}\}'  # 特别添加对{{变量}}的标准格式支持
        for pattern in [standard_pattern] + self.placeholder_patterns:
            pattern_search = re.escape(pattern) if isinstance(pattern, str) and not pattern.startswith(r'\{\{') else pattern
            if re.search(pattern_search, template_text):
                placeholder_matches.append((re.escape(pattern), pattern))
        
        # 如果没有找到占位符匹配，返回空字符串
        if not placeholder_matches:
            return ""
            
        # 遍历所有匹配的占位符格式
        for escaped_pattern, raw_pattern in placeholder_matches:
            # 情况1: 模板中格式为"标签: 占位符"
            # 在目标文档中应该是"标签: 实际值"
            label_pattern = r'([^:：]+[:：]\s*)' + escaped_pattern
            label_match = re.search(label_pattern, template_text)
            if label_match:
                label = label_match.group(1).strip()
                # 在目标文本中查找相同标签后面的值
                value_pattern = re.escape(label) + r'(.*?)(?:$|\n|,|，|;|；|\s\s)'
                value_match = re.search(value_pattern, target_text)
                if value_match:
                    return value_match.group(1).strip()
            
            # 情况2: 占位符在模板中的位置替换
            # 分析模板文本中占位符的前后文本
            parts = re.split(escaped_pattern, template_text, 1)  # 限制分割一次，增加精确度
            if len(parts) == 2:
                prefix = parts[0]
                suffix = parts[1]
                
                # 如果前缀和后缀都存在于目标文本中
                if prefix in target_text and suffix in target_text:
                    # 提取前缀和后缀之间的文本作为值
                    pattern = re.escape(prefix) + r'(.*?)' + re.escape(suffix)
                    match = re.search(pattern, target_text)
                    if match:
                        return match.group(1).strip()
                
                # 如果只有前缀存在
                elif prefix in target_text and prefix.strip():
                    parts_after_prefix = target_text.split(prefix, 1)
                    if len(parts_after_prefix) > 1:
                        # 提取前缀后的文本，直到下一个自然分隔符
                        after_prefix = parts_after_prefix[1].strip()
                        # 查找自然终止符
                        end_match = re.search(r'[。.;；,，\n]|\s{2,}', after_prefix)
                        if end_match:
                            return after_prefix[:end_match.start()].strip()
                        else:
                            # 如果没有自然终止符，使用固定字符数
                            return after_prefix[:50].strip() if len(after_prefix) > 50 else after_prefix
            
            # 情况3: 使用简单的位置替换
            if raw_pattern == standard_pattern:  # 特殊处理标准{{变量}}格式
                # 检查是否是表格或段落中的单独项
                if template_text.strip() == '{{' + placeholder + '}}':
                    return target_text.strip()
                
                # 检查标准标签格式，如"姓名：{{姓名}}"
                label_pattern = r'([^:：]+)[:：]\s*\{\{' + re.escape(placeholder) + r'\}\}'
                label_match = re.search(label_pattern, template_text)
                if label_match:
                    label = label_match.group(1).strip()
                    value_pattern = re.escape(label) + r'[:：]\s*(.*?)(?:$|\n|,|，|;|；|\s\s)'
                    value_match = re.search(value_pattern, target_text)
                    if value_match:
                        return value_match.group(1).strip()
        
        return ""

    def _extract_from_text(self, text, data, placeholders):
        """
        从文本中提取数据
        
        Args:
            text: 要提取的文本
            data: 数据存储字典
            placeholders: 占位符列表
        """
        for placeholder in placeholders:
            # 查找该占位符所有可能的表示形式
            for pattern in self.placeholder_patterns:
                try:
                    # 将占位符名称替换到模式中
                    search_pattern = pattern.replace("(.*?)", re.escape(placeholder))
                    if search_pattern in text:
                        value = text.replace(search_pattern, "").strip()
                        if value and placeholder in data:
                            data[placeholder].append(value)
                            logger.debug(f"从文本提取到 {placeholder}: {value}")
                except Exception as e:
                    logger.error(f"从文本提取 {placeholder} 时出错: {str(e)}")
    
    def export_to_excel(self, data, output_path):
        """
        将提取的数据导出到Excel
        
        Args:
            data: pandas.DataFrame 数据
            output_path: 输出Excel文件路径
        """
        try:
            data.to_excel(output_path, index=False)
            logger.info(f"数据已成功导出到 {output_path}")
            return True
        except Exception as e:
            logger.error(f"导出数据到Excel时出错: {str(e)}")
            raise

    def _extract_image_from_document(self, placeholder):
        """
        从文档中提取图片
        
        Args:
            placeholder: 图片占位符名称
            
        Returns:
            str: 提取的图片路径，如果未找到则返回空字符串
        """
        if not hasattr(self, 'current_doc_path') or not self.current_doc_path:
            logger.error("未设置当前文档路径，无法提取图片")
            return ""
            
        try:
            # 提取图片字段名（例如从"img:姓名"提取出"姓名"）
            field_name = placeholder
            if "img:" in placeholder:
                # 处理不同格式的占位符
                if placeholder.startswith("img:"):
                    field_name = placeholder.split("img:")[1]
                elif "img:" in placeholder:
                    field_name = placeholder.split("img:")[1]
                    
                # 清理字段名中可能的额外字符（例如从"{img:姓名}}"中的"姓名}}"提取"姓名"）
                field_name = re.sub(r'[{}]', '', field_name)
                field_name = field_name.strip()
                
                logger.debug(f"从占位符 {placeholder} 中提取字段名: {field_name}")
            
            # 获取已提取的数据值（如果有）
            data_value = None
            if hasattr(self, 'current_extracted_data') and self.current_extracted_data:
                # 尝试通过字段名获取数据
                if field_name in self.current_extracted_data:
                    data_value = str(self.current_extracted_data[field_name])
                    data_value = re.sub(r'[\\/*?:"<>|]', '_', data_value)  # 替换文件名中的非法字符
                    logger.debug(f"使用提取的数据值 '{data_value}' 命名图片")
                
            # 创建文档对象
            from docx import Document
            doc = Document(self.current_doc_path)
            
            # 创建图片保存目录
            doc_dir = os.path.dirname(self.current_doc_path)
            doc_name = os.path.splitext(os.path.basename(self.current_doc_path))[0]
            images_dir = os.path.join(doc_dir, f"{doc_name}_images")
            if not os.path.exists(images_dir):
                os.makedirs(images_dir)
                
            # 提取文档中的所有图片
            image_index = 1
            image_paths = []
            
            # 处理文档关系部分中的图片
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    # 保存图片到本地
                    image_data = rel.target_part.blob
                    if not image_data:
                        continue
                    
                    # 使用提取的数据值或字段名构建文件名
                    if data_value:
                        filename = f"{data_value}"
                    else:
                        filename = f"{field_name}"
                    
                    # 限制文件名长度并添加索引
                    if len(filename) > 50:
                        filename = filename[:50]
                    filename = f"{filename}_{image_index}.png"
                    #姓名_袁修平_1.png     filename = f"{filename}_{image_index}.png"
                    image_path = os.path.join(images_dir, filename)
                    with open(image_path, "wb") as f:
                        f.write(image_data)
                    
                    image_paths.append(image_path)
                    image_index += 1
                    logger.debug(f"已提取图片到: {image_path}")
            
            # 如果找到了图片，返回第一张图片的路径
            if image_paths:
                return image_paths[0]
                
            logger.warning(f"未在文档中找到任何图片: {self.current_doc_path}")
            return ""
        except Exception as e:
            logger.error(f"提取图片时出错: {str(e)}")
            return ""

if __name__ == "__main__":
    # 简单的测试功能
    print("Word数据提取器 - 测试模式")
    
    # 创建提取器实例
    extractor = WordExtractor(None)
    
    # 检查命令行参数
    if len(sys.argv) > 2:
        template_path = sys.argv[1]
        target_path = sys.argv[2]
        
        print(f"使用模板: {template_path}")
        print(f"目标文档: {target_path}")
        
        try:
            # 检测占位符
            extractor.template_doc = Document(template_path)
            placeholders = extractor.detect_placeholders(extractor.template_doc)
            print(f"检测到的占位符: {placeholders}")
            
            # 提取数据
            data = extractor.extract_data(target_path)
            print("\n提取的数据:")
            for key, value in data.items():
                print(f"  {key}: {value}")
                
        except Exception as e:
            print(f"错误: {str(e)}")
    else:
        print("用法: python word_extractor.py <模板路径> <目标文档路径>")
        print("示例: python word_extractor.py template.docx target.docx") 