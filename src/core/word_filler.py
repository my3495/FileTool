"""
Word数据填充模块
将Excel数据填充到Word文档
"""
import os
import re
from copy import deepcopy
import io

import pandas as pd
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches
from docx.table import _Cell
from docx.oxml.ns import qn
from lxml import etree
from loguru import logger
from src.core.placeholder_parser import BasePlaceholderParser, PlaceholderParserFactory


class WordFiller(BasePlaceholderParser):
    """
    Word数据填充类
    将Excel中的数据填充到Word模板中
    """
    
    def __init__(self, config_manager=None):
        """
        初始化填充器
        
        Args:
            config_manager: 配置管理器，可选
        """
        super().__init__(config_manager)
        # 设置默认图片宽度
        self.default_image_width = 4.0  # 英寸
        if config_manager:
            self.default_image_width = config_manager.get("default_image_width", 4.0)
        
        # 设置占位符模式
        self.placeholder_patterns = [
            r'\{\{(.*?)\}\}',  # {{变量}}
            r'\$\{(.*?)\}',    # ${变量}
            r'#(.*?)#',        # #变量#
            r'\{\{img:(.*?)\}\}'  # {{img:变量}}
        ]
        
        # 创建解析器工厂
        self.parser_factory = PlaceholderParserFactory(config_manager)
            
        logger.info("初始化Word数据填充器")
    
    def load_excel_data(self, excel_path, sheet_name=0):
        """
        加载Excel数据
        
        Args:
            excel_path: Excel文件路径
            sheet_name: 工作表名称或索引
            
        Returns:
            pandas.DataFrame: 加载的数据
        """
        try:
            df = pd.read_excel(excel_path, sheet_name=sheet_name)
            logger.info(f"成功从Excel加载数据，共{len(df)}行")
            return df
        except Exception as e:
            logger.error(f"加载Excel数据时出错: {str(e)}")
            raise
    
    def fill_template(self, template_path, data_row, output_path):
        """
        使用数据行填充Word模板
        
        Args:
            template_path: Word模板文件路径
            data_row: 包含数据的Series或dict
            output_path: 输出文件路径
            
        Returns:
            bool: 操作是否成功
        """
        try:
            doc = Document(template_path)
            
            # 替换段落、表格、页眉和页脚中的占位符
            self._replace_all_placeholders(doc, data_row)
            
            # 保存文档
            doc.save(output_path)
            logger.info(f"文档已保存到 {output_path}")
            return True
        except Exception as e:
            logger.error(f"填充模板时出错: {str(e)}")
            raise
    
    def _replace_all_placeholders(self, doc, data_row):
        """替换文档中所有地方的占位符"""
        # 段落
        for para in doc.paragraphs:
            self._replace_text(para, data_row)
        
        # 表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_text(para, data_row, parent=cell)

        # 页眉和页脚
        for section in doc.sections:
            # 页眉
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for para in header.paragraphs:
                        self._replace_text(para, data_row)
                    for table in header.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    self._replace_text(para, data_row, parent=cell)
            # 页脚
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for para in footer.paragraphs:
                        self._replace_text(para, data_row)
                    for table in footer.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    self._replace_text(para, data_row, parent=cell)

    def _replace_text(self, para, data, parent=None):
        """
        替换一个段落中的所有占位符。
        这个新方法更简单、更健壮。
        """
        if not para.text:
            return
            
        # 查找此段落中的所有占位符
        placeholders = []
        for parser in self.parser_factory.parsers:
            # 使用 para.text 获取最新文本状态
            placeholders.extend(parser.find_all_placeholders(para.text))

        # 按占位符在文本中的出现顺序排序，从后往前替换，避免索引失效
        placeholders.sort(key=lambda p: para.text.find(p[1]), reverse=True)

        for p_name, p_full, _ in placeholders:
            if self.is_image_placeholder(p_name):
                field_name = self.get_image_field_name(p_name)
                img_path = str(data.get(field_name, ''))
                if os.path.exists(img_path):
                    self._replace_image_in_paragraph(para, p_full, img_path, parent)
                else:
                    logger.warning(f"图片文件不存在: {img_path}")
                    # 如果图片不存在，替换为提示信息
                    self.inline_replace(para, p_full, f'[图片丢失: {img_path}]')
            elif p_name in data:
                value = str(data[p_name])
                self.inline_replace(para, p_full, value)

    def inline_replace(self, para, placeholder, value):
        """
        在段落中进行内联替换，以保留格式。
        """
        # 找到占位符所在的 runs
        start_run_idx, start_char_idx, end_run_idx, end_char_idx = self._find_placeholder_runs(para, placeholder)

        if start_run_idx is None:
            # 如果在 run 级别找不到，尝试在段落级别进行简单替换（会丢失格式）
            if placeholder in para.text:
                 para.text = para.text.replace(placeholder, value)
            return

        # 进行替换
        start_run = para.runs[start_run_idx]
        
        if start_run_idx == end_run_idx:
            # 占位符在同一个 run 中
            before = start_run.text[:start_char_idx]
            after = start_run.text[end_char_idx:]
            start_run.text = before + value + after
        else:
            # 占位符跨越多个 run
            # 1. 修改第一个 run
            start_run.text = start_run.text[:start_char_idx] + value

            # 2. 清空中间的 runs
            for i in range(start_run_idx + 1, end_run_idx):
                # 安全地访问 run，并清除其文本
                if i < len(para.runs):
                    para.runs[i].text = ''

            # 3. 修改最后一个 run
            if end_run_idx < len(para.runs):
                end_run = para.runs[end_run_idx]
                end_run.text = end_run.text[end_char_idx:]

    def _replace_image_in_paragraph(self, para, placeholder, img_path, parent):
        """替换段落中的图片占位符"""
        # 根据上下文（是否在表格中）确定图片宽度
        width = None
        if isinstance(parent, _Cell):
            width = self._get_image_width_for_cell(parent)
        
        # 如果无法确定宽度或不在表格中，则使用默认宽度
        if not width:
            width = Inches(self.default_image_width)

        # 先清除占位符文本
        self.inline_replace(para, placeholder, '')
        
        # 在段落末尾添加图片
        try:
            para.add_run().add_picture(img_path, width=width)
            logger.debug(f"成功插入图片 '{img_path}'，宽度: {width.inches:.2f} 英寸")
        except Exception as e:
            logger.error(f"插入图片 '{img_path}' 时出错: {e}")

    def _get_image_width_for_cell(self, cell):
        """计算并返回适合单元格的图片宽度（英寸）"""
        width = None
        try:
            # 尝试直接获取单元格宽度
            if cell.width:
                # 留出一些边距 (e.g., 90% of cell width)
                width = Inches(cell.width.inches * 0.9)
                logger.debug(f"单元格直接宽度: {cell.width.inches:.2f} 英寸, 图片宽度: {width.inches:.2f} 英寸")
                return width
        except Exception:
             # 在某些情况下 .width 会失败
             pass

        try:
            # 备选方案: 通过表格的列宽来推断
            table = cell._element.getparent().getparent() # 获取父级<table>元素
            grid = table.find('.//w:tblGrid', namespaces=table.nsmap)
            if grid is not None:
                cols = grid.findall('.//w:gridCol', namespaces=grid.nsmap)
                
                # 找到当前单元格属于第几列
                cell_idx = cell._element.getparent().index(cell._element)
                
                if cell_idx < len(cols):
                    col_width_twips = cols[cell_idx].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w')
                    if col_width_twips:
                        # 转换 twips 为英寸 (1 英寸 = 1440 twips)
                        # 留出一些边距
                        width = Inches(int(col_width_twips) / 1440 * 0.9)
                        logger.debug(f"推断的列宽: {int(col_width_twips) / 1440:.2f} 英寸, 图片宽度: {width.inches:.2f} 英寸")
                        return width
        except Exception as e:
            logger.warning(f"无法通过列宽推断图片大小: {e}")

        # 如果所有方法都失败了，返回None
        logger.warning("无法确定单元格宽度，将使用默认图片宽度。")
        return None

    def _find_placeholder_runs(self, para, placeholder):
        """找到占位符在段落 runs 中的起始和结束位置"""
        runs_text = [r.text for r in para.runs]
        full_text = "".join(runs_text)
        
        p_start = full_text.find(placeholder)
        if p_start == -1:
            return None, None, None, None
        p_end = p_start + len(placeholder)

        # 计算起始 run 和字符索引
        current_len = 0
        start_run_idx, start_char_idx = -1, -1
        for i, run_text in enumerate(runs_text):
            if current_len <= p_start < current_len + len(run_text):
                start_run_idx = i
                start_char_idx = p_start - current_len
                break
            current_len += len(run_text)

        # 计算结束 run 和字符索引
        current_len = 0
        end_run_idx, end_char_idx = -1, -1
        for i, run_text in enumerate(runs_text):
            if current_len < p_end <= current_len + len(run_text):
                end_run_idx = i
                end_char_idx = p_end - current_len
                break
            current_len += len(run_text)
            
        return start_run_idx, start_char_idx, end_run_idx, end_char_idx
    
    def _append_doc(self, source_doc, target_doc):
        """将 source_doc 的内容附加到 target_doc，并正确处理图片"""
        # --- 最终修复方案：解构并重构图片 ---
        # 1. 创建一个从旧 rId 到新 rId 的映射
        rid_map = {}
        for rel in source_doc.part.rels:
            if source_doc.part.rels[rel].reltype == RT.IMAGE:
                # a. 获取源文档中的图片部件
                image_part = source_doc.part.rels[rel].target_part
                
                # b. 从图片部件中提取二进制数据
                image_blob = image_part.blob
                
                # c. 在目标文档中，通过 add_picture 注册图片并获取新的 rId。
                #    这是确保图片被正确添加到目标包的关键。
                #    我们将图片添加到一个临时段落，然后立即删除该段落，
                #    目的仅仅是为了获得新图片的关系ID。
                temp_para = target_doc.add_paragraph()
                run = temp_para.add_run()
                inline_shape = run.add_picture(io.BytesIO(image_blob))
                new_rid = inline_shape._inline.graphic.graphicData.pic.blipFill.blip.embed
                
                # d. 清理临时创建的段落
                p = temp_para._element
                p.getparent().remove(p)

                # e. 存储rId映射
                rid_map[rel] = new_rid

        # 2. 复制并附加内容
        for element in source_doc.element.body:
            # 使用 lxml 的 tostring/fromstring 确保创建完全独立的元素副本
            new_element = etree.fromstring(etree.tostring(element))
            
            # 3. 在新元素中查找所有图片引用 (blip)，并更新 rId
            for blip in new_element.iter(qn('a:blip')):
                r_embed = blip.get(qn('r:embed'))
                if r_embed in rid_map:
                    # 更新 rId
                    blip.set(qn('r:embed'), rid_map[r_embed])

            target_doc.element.body.append(new_element)

    def batch_fill_templates(self, template_path, excel_path, output_dir, filename_pattern=None, merge_output=False):
        """
        批量填充模板
        
        Args:
            template_path: Word模板文件路径
            excel_path: Excel数据文件路径
            output_dir: 输出目录
            filename_pattern: 文件名模式，可包含列名如 "{姓名}_{部门}"
            merge_output: 是否合并输出到单个文件
            
        Returns:
            list: 生成的文件路径列表
        """
        try:
            os.makedirs(output_dir, exist_ok=True)
            df = self.load_excel_data(excel_path)
            if df.empty:
                logger.warning("Excel文件中没有数据")
                return []

            output_files = []

            if merge_output:
                # --- 合并输出模式 ---
                logger.info("开始执行合并输出模式...")
                main_doc = None
                
                for index, row in df.iterrows():
                    logger.debug(f"正在处理第 {index + 1} 行数据用于合并...")
                    row_dict = row.to_dict()
                    
                    # 1. 为当前行数据在内存中创建一个独立的、填充好的文档
                    row_doc = Document(template_path)
                    self._replace_all_placeholders(row_doc, row_dict)
                    
                    # 2. 将这个新文档附加到主文档
                    if main_doc is None:
                        # 第一个文档，直接将其作为主文档
                        main_doc = row_doc
                    else:
                        # 后续文档，先添加分页符再附加内容
                        main_doc.add_page_break()
                        self._append_doc(row_doc, main_doc)
                
                # 循环结束后，保存主文档
                if main_doc:
                    merged_filename = "合并文档.docx"
                    # 检查是否有自定义文件名模式，如果有，则使用第一行数据生成
                    if filename_pattern:
                        first_row_dict = df.iloc[0].to_dict()
                        temp_filename = filename_pattern
                        for col, value in first_row_dict.items():
                            temp_filename = temp_filename.replace(f"{{{col}}}", str(value))
                        
                        temp_filename = temp_filename.replace("{序号}", "1")
                        
                        # 移除路径中的非法字符
                        safe_filename = re.sub(r'[\\/*?:"<>|]', "_", temp_filename)
                        merged_filename = f"{safe_filename}_合并.docx"
                    
                    merged_path = os.path.join(output_dir, merged_filename)
                    main_doc.save(merged_path)
                    output_files.append(merged_path)
                    logger.info(f"合并文档已保存到: {merged_path}")

            else:
                # --- 分开保存模式 (优化文件名生成) ---
                logger.info("开始执行分开保存模式...")
                for index, row in df.iterrows():
                    logger.debug(f"正在处理第 {index + 1} 行数据用于分开保存...")
                    row_dict = row.to_dict()
                    
                    # 生成文件名
                    filename = f"文档_{index + 1}.docx"  # 默认文件名
                    if filename_pattern:
                        temp_filename = filename_pattern
                        # 替换列名占位符
                        for col, value in row_dict.items():
                            temp_filename = temp_filename.replace(f"{{{col}}}", str(value))
                        
                        # 替换特殊占位符 {序号}
                        temp_filename = temp_filename.replace("{序号}", str(index + 1))

                        # 移除路径中的非法字符并添加后缀
                        safe_filename = re.sub(r'[\\/*?:"<>|]', "_", temp_filename)
                        filename = f"{safe_filename}.docx"
                    
                    output_path = os.path.join(output_dir, filename)
                    
                    # 填充并保存单个文件
                    self.fill_template(template_path, row_dict, output_path)
                    output_files.append(output_path)

            logger.info(f"批量生成完成，共生成 {len(output_files)} 个文件")
            return output_files
        except Exception as e:
            logger.error(f"批量填充模板时出错: {str(e)}")
            raise
    
    def is_image_placeholder(self, placeholder_name):
        """
        检查占位符是否为图片占位符
        
        Args:
            placeholder_name: 占位符名称
            
        Returns:
            bool: 如果是图片占位符则为True，否则为False
        """
        return placeholder_name.startswith("img:") or "img:" in placeholder_name
        
    def get_image_field_name(self, placeholder_name):
        """
        从图片占位符中获取字段名
        
        Args:
            placeholder_name: 占位符名称
            
        Returns:
            str: 字段名
        """
        if placeholder_name.startswith("img:"):
            return placeholder_name
        elif "img:" in placeholder_name:
            return placeholder_name.split("img:")[1]
        return placeholder_name 